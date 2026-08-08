#!/usr/bin/env python3
"""Generate ConsTrakr Android rebuild specification DOCX."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date


def add_title_page(doc: Document) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("ConsTrakr Ecosystem\nFeature Inventory &\nAndroid Rebuild Specification")
    run.bold = True
    run.font.size = Pt(24)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Version 1.0 — {date.today().strftime('%B %d, %Y')}")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Covers IMS People (web), ConsTrakr iOS app features, sync API, "
        "and a phased plan to recreate the field app on Android."
    )
    nr.font.size = Pt(11)
    doc.add_page_break()


def add_mermaid_block(doc: Document, title: str, code: str) -> None:
    doc.add_heading(title, level=3)
    p = doc.add_paragraph()
    p.style = "Intense Quote"
    run = p.add_run(code.strip())
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    doc.add_paragraph(
        "Paste the block above into https://mermaid.live to render the diagram.",
        style="Intense Quote",
    )


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_title_page(doc)

    # ── TOC overview ──
    doc.add_heading("Document Contents", level=1)
    for item in [
        "1. System Overview",
        "2. Web (People / IMS) Features",
        "3. ConsTrakr App Features (iOS reference)",
        "4. Constrakr Sync API",
        "5. Android Rebuild — Phase by Phase",
        "6. Architecture Diagrams (Mermaid)",
        "7. Web-Only vs App Scope",
        "8. Android Technology Mapping",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ── 1. Overview ──
    doc.add_heading("1. System Overview", level=1)
    doc.add_paragraph(
        "ConsTrakr is an offline-first attendance and roster system. "
        "The IMS People module (web) is the admin back-office. "
        "The ConsTrakr mobile app is the field device for face-based Time In/Out. "
        "Both sides sync through the Constrakr REST API at /constrakr-api/."
    )
    add_table(doc, ["Layer", "Role"], [
        ["Web (People)", "Admin: roster, DTR, payroll, reports, device assignment, void passcode"],
        ["App (ConsTrakr)", "Field: face attendance, offline roster, job sites, sync"],
        ["API", "Sync bridge; JWT auth separate from IMS session login"],
    ])
    doc.add_paragraph(
        "Shared security: a user's 6-digit void passcode (Profile → Edit Profile on web) "
        "is also the device admin code for protected app actions (default site, GPS, employee edit)."
    )

    doc.add_page_break()

    # ── 2. Web ──
    doc.add_heading("2. Web (People / IMS) Features", level=1)
    doc.add_paragraph("Base path: /people/  |  Access: can_view_people_pages  |  Data scoped by user group")

    doc.add_heading("2.1 Navigation", level=2)
    add_table(doc, ["Tab", "Purpose"], [
        ["Home", "Dashboard KPIs"],
        ["Employees", "Roster management"],
        ["Time & attendance", "DTR / punches"],
        ["Job sites", "Sites + payroll per site"],
        ["Devices", "ConsTrakr phones/tablets"],
        ["Departments", "Dept/position catalog"],
        ["Payroll & reports", "Payroll hub + exports"],
        ["Support", "Sync health & troubleshooting"],
    ])

    doc.add_heading("2.2 Employees", level=2)
    for b in [
        "List: search, filters (enrolled, status, job site), pagination, enrollment thumbnails",
        "Detail: profile, enrollment photos/embeddings, ID doc, recent attendance, leaves, sync times",
        "Edit: name, dept/position, job site, pay fields, govt IDs, work weekdays, break",
        "Remove from web: soft-delete (void passcode required); history retained",
        "Restore to app: clears soft-delete",
        "Payroll setup: allowances, recurring deductions, adjustments, per-site rate overrides",
        "Email payslip",
        "Support lookup by server ID, local ID, or employee code",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("2.3 Time & Attendance (DTR)", level=2)
    for b in [
        "Weekly DTR grid with filters (week, employee, site, check type, void status, search)",
        "Void single punch — 6-digit void passcode required",
        "Void full day — void passcode required",
        "Manual DTR correction — voids day, creates manual in/out (no passcode)",
        "Attendance detail view",
        "Voided punches excluded from default views; filterable",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("2.4 Job Sites", level=2)
    for b in [
        "CRUD: name, location, GPS, radius",
        "Payroll standard: PH Labor Code or custom; OT/holiday multipliers, pay frequency",
        "Per-site: Manpower, Period/Weekly DTR, Payroll run, Settings",
        "Delete site: soft-delete + unassign employees (void passcode)",
        "CSV exports on DTR and payroll",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("2.5 Devices", level=2)
    for b in [
        "List registered ConsTrakr devices (name, local ID, app version, last seen, Active/Blocked status)",
        "Assign one or more IMS users — any assigned user's 6-digit admin code unlocks device",
        "Admin code readiness indicator per user",
        "Edit device name",
        "Remote access: Block device (optional reason) — stops scan, CRUD, and sync on app",
        "Unblock device — restores app access after Check again on device",
        "Remove from list (soft-delete); blocked devices must be unblocked first",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("2.6 Remote Device Blocking (SaaS)", level=2)
    for b in [
        "Web: People → Devices → Edit → Remote access → Block / Unblock",
        "Optional block reason shown on the app lock screen",
        "API: is_blocked, blocked_reason, blocked_at returned on GET/POST /devices",
        "API: all sync write endpoints reject blocked devices (HTTP 403 device_blocked)",
        "App sends X-Device-Local-Id header on authenticated API requests",
        "POST /devices still works as heartbeat when blocked (does not auto-unblock)",
        "Offline: app keeps last known block state until next server check",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("2.7 Departments", level=2)
    doc.add_paragraph(
        "Manage department categories and positions per user group. Add, rename, soft-delete. Seeded defaults."
    )

    doc.add_heading("2.8 Payroll & Reports", level=2)
    for b in [
        "Reports hub: payslips, monthly summary, consolidated payroll, bank disbursement",
        "Attendance exceptions, leave records, OT approvals, holiday calendar",
        "Government remittance (SSS, PhilHealth, Pag-IBIG, tax), 13th month",
        "Payroll period lock; recurring loan deductions on lock",
        "Payslip view/print/email; CSV exports",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("2.9 Support Tools", level=2)
    add_table(doc, ["Tool", "Purpose"], [
        ["Sync health hub", "Overview"],
        ["Punch check", "Lookup punches by ID/code/date"],
        ["Enrollment gaps", "Missing face templates/photos/ID"],
        ["Sync integrity", "Removed, stale, duplicate employees/sites"],
        ["Batch employee check", "Bulk code verification"],
        ["Job site diagnostics", "Duplicates, invalid assignments"],
        ["API health", "Health endpoint + clock drift"],
        ["Help", "Troubleshooting guide"],
    ])

    doc.add_heading("2.10 Profile / Void Passcode", level=2)
    doc.add_paragraph(
        "Exactly 6-digit void passcode under Profile → Edit Profile. Used for web void actions "
        "and as device admin code when user is assigned to a device."
    )

    doc.add_page_break()

    # ── 3. App ──
    doc.add_heading("3. ConsTrakr App Features (iOS Reference)", level=1)
    doc.add_paragraph("Offline-first field app. Default job site drives Dashboard, Employees, DTR, and scanner.")

    doc.add_heading("3.1 Tab Structure", level=2)
    add_table(doc, ["Tab", "Screen", "Purpose"], [
        ["Dashboard", "Coverage + roster stats", "Today's attendance health"],
        ["Employees", "Filtered roster", "Register, view, edit, delete"],
        ["Scanner", "Face attendance", "Time In / Time Out"],
        ["Time Record", "DTR grid", "Daily in/out per employee"],
        ["More", "Job sites + Settings", "Sites, sync, config"],
    ])

    doc.add_heading("3.2 Dashboard", level=2)
    for b in [
        "Scoped to default job site only",
        "Coverage ring: full day=1.0, in-only=0.5, absent=0",
        "Chips: Today in/out, Absent, Incomplete",
        "Default site card → tap opens DTR tab",
        "Needs attention when coverage < 70% or absent > 0",
        "Tiles: Enrolled, Unassigned, Checked in, Incomplete, Pending sync",
        "Pull-to-sync (quick, all scope); online/offline indicator",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.3 Employees", level=2)
    doc.add_paragraph("List filtered to default job site. Search by name/code/dept. Swipe delete. Register toolbar.")
    doc.add_paragraph("Registration wizard (3 steps):")
    for b in [
        "Step 1: Details — name, dept/position, job site, auto code YYYYMM+seq",
        "Step 2: Government ID — type, number, camera scan (skippable)",
        "Step 3: Face scan — blink, optional 3D depth, 5 poses (center/left/right/up/down)",
    ]:
        doc.add_paragraph(b, style="List Bullet 2")
    doc.add_paragraph(
        "Edit: admin code gate (6-digit). Editable: name, dept, position, site. "
        "Delete: local wipe + server soft-delete queue."
    )

    doc.add_heading("3.4 Attendance Scanner", level=2)
    doc.add_paragraph("Pre-scan gates: geofence → Time In/Out confirm → supervisor PIN (optional) → GPS per employee site")
    doc.add_paragraph("Pipeline: warmup → liveness (blink + poses + 3D) → anti-spoof → AdaFace match → clock integrity → site gate → duplicate check")
    doc.add_paragraph("Outcomes: recognized (green), already recorded (orange), wrong site/clock (red), unknown (auto-cancel 3s)")

    doc.add_heading("3.5 DTR (Time Record)", level=2)
    for b in [
        "Requires default job site; date picker",
        "One row per employee: earliest In, latest Out",
        "Corrected badge for manual server corrections; punch photo thumbnails",
        "Pull-to-sync: attendance scope for selected date",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.6 Job Sites", level=2)
    for b in [
        "Local JSON storage (not main DB)",
        "List, add, edit, swipe delete, pending sync banner",
        "Editor: name, location, map pin, radius 30–100m, set default",
        "Admin code required for save and set-default",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.7 Settings", level=2)
    add_table(doc, ["Section", "Features"], [
        ["Appearance", "System / Light / Dark theme"],
        ["Status", "Network, pending uploads, last sync, progress"],
        ["Sync", "Auto sync, interval 3–60 min, Wi-Fi-only uploads, full sync"],
        ["Sync account", "sync_admin JWT sign-in/out"],
        ["Scanner", "Match threshold, liveness presets, per-pose toggles"],
        ["Job sites & GPS", "Geofence toggle + default site (admin code gates)"],
        ["Supervisor PIN", "Optional 4–12 digit PIN before punch (local hash)"],
        ["Advanced", "Device name/ID, server URL, restore, diagnostics"],
    ])

    doc.add_heading("3.8 Admin Code Gates", level=2)
    for b in [
        "Exactly 6-digit, server-verified via POST /devices/verify-admin-code",
        "iPhone-style 6-dot passcode keypad (auto-submit on 6th digit)",
        "Required: change default site, toggle require GPS, save/set-default job site, edit employee",
        "Requires one or more assigned users on device with at least one passcode set on web",
        "Any assigned user's code works (multi-user device assignment on web)",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.9 Remote Device Blocking (SaaS)", level=2)
    for b in [
        "Full-screen DeviceBlockedView when is_blocked is true",
        "Blocks: attendance scanner, employee register/edit, sync uploads",
        "Shows admin block reason, device ID, and Check again button",
        "Checks on launch, app resume, and every sync attempt",
        "Cached locally in UserDefaults — enforced offline until server clears block",
        "APIService sends X-Device-Local-Id on all authenticated requests",
        "HTTP 403 device_blocked from API also triggers local lock",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.10 Sync System", level=2)
    add_table(doc, ["Mode", "Use"], [
        ["quick", "Pull-to-sync, foreground refresh"],
        ["full", "Full roster + child asset reconcile"],
    ])
    add_table(doc, ["Scope", "Use"], [
        ["all", "Everything"],
        ["employees", "Roster + embeddings/photos/ID"],
        ["attendance", "Punches only (DTR date focus)"],
    ])
    doc.add_paragraph(
        "Push order: device registration → job sites → departments → pending deletions → "
        "pull roster → push employees → embeddings → photos → ID docs → attendance → reconcile voids → pull attendance."
    )
    doc.add_paragraph("Offline-first: write locally as pending; sync when online + authenticated.")

    doc.add_page_break()

    # ── 4. API ──
    doc.add_heading("4. Constrakr Sync API", level=1)
    doc.add_paragraph("Base: {host}/constrakr-api  |  Auth: Bearer JWT (except health + login)")
    add_table(doc, ["Method", "Path", "Purpose"], [
        ["GET", "/health", "Health + server time"],
        ["POST", "/auth/admin/login", "JWT (sync_admin)"],
        ["GET/POST", "/employees", "Roster sync (delta: updated_since)"],
        ["GET/POST", "/employees/check", "Existence check"],
        ["PUT/DELETE", "/employees/{id}", "Update / soft-delete"],
        ["POST", "/employees/{id}/restore", "Restore soft-deleted"],
        ["GET/POST", "/job-sites", "Site catalog"],
        ["PUT/DELETE", "/job-sites/{id}", "Update / soft-delete"],
        ["GET/POST", "/devices", "Device registration"],
        ["POST", "/devices/verify-admin-code", "6-digit admin code"],
        ["GET", "/departments", "Department catalog"],
        ["GET/POST", "/face-embeddings", "Encrypted face templates (5 poses)"],
        ["GET/POST", "/face-enrollment-photos", "Enrollment JPEGs"],
        ["GET/POST", "/employee-id-documents", "Govt ID photo + metadata"],
        ["GET/POST", "/attendance", "Punch sync"],
    ])
    doc.add_paragraph("Idempotency: every entity has local_id (UUID) separate from server_id.")
    doc.add_paragraph(
        "Device header: authenticated requests include X-Device-Local-Id for SaaS remote blocking. "
        "Device JSON includes is_blocked, blocked_reason, blocked_at."
    )

    doc.add_page_break()

    # ── 5. Phases ──
    doc.add_heading("5. Android Rebuild — Phase by Phase", level=1)
    doc.add_paragraph(
        "Build in this order. Each phase lists deliverables, screens, API calls, and local storage."
    )

    phases = [
        {
            "title": "Phase 1 — Foundation & Project Setup",
            "goal": "Runnable app shell with networking, auth, and database scaffolding.",
            "deliverables": [
                "Kotlin + Jetpack Compose project (min SDK 26+ recommended)",
                "Navigation: BottomNav with 5 tabs (placeholder screens)",
                "Room database: Employee, Attendance, FaceEmbedding, FaceEnrollmentPhoto, EmployeeIdDocument entities",
                "EncryptedSharedPreferences / Keystore for JWT",
                "Retrofit/Ktor client: base URL + /constrakr-api prefix",
                "NetworkMonitor (ConnectivityManager callback)",
                "App theme: Material 3, System/Light/Dark",
            ],
            "api": ["GET /health", "POST /auth/admin/login"],
            "storage": ["Room schema v1", "Encrypted prefs: token, username, expiry"],
            "exit": "App launches, tabs navigate, can ping health and store JWT after login",
        },
        {
            "title": "Phase 2 — Settings & Sync Account",
            "goal": "Admin sign-in, server URL config, sync status UI.",
            "deliverables": [
                "Settings screen (List/inset grouped style)",
                "Sync account section: username/password, sign in/out",
                "Advanced: server URL field, Test API (health + login)",
                "Status section: online/offline, last sync, pending count placeholder",
                "About: app name, face engine label, developer credit",
                "SyncQueue singleton: isSyncing, lastError, lastSyncDate observables",
            ],
            "api": ["GET /health", "POST /auth/admin/login"],
            "storage": ["settings.apiBaseURL", "SyncAuthStore equivalent"],
            "exit": "User can sign in as sync_admin and see connection status",
        },
        {
            "title": "Phase 3 — Job Sites (Local + Sync)",
            "goal": "Job site CRUD, default site, sync to server.",
            "deliverables": [
                "JobSiteStore (SharedPreferences JSON + pending upload/delete queues)",
                "More → Job Sites list, add/edit, swipe delete",
                "Map pin editor (Google Maps or OSM)",
                "Settings → Job Sites & GPS: default site picker, geofence toggle (UI only; gate in Phase 8)",
                "Push/pull job sites on sync",
                "Pull-to-sync on job sites list",
            ],
            "api": ["GET/POST /job-sites", "PUT/DELETE /job-sites/{id}"],
            "storage": ["jobSitesJSON", "defaultJobSiteId", "pendingJobSiteUploads/Deletions"],
            "exit": "Sites persist offline, sync when signed in, default site drives later phases",
        },
        {
            "title": "Phase 4 — Employee Roster (No Face Yet)",
            "goal": "List, register details, edit profile, delete — without scanner.",
            "deliverables": [
                "Employee list filtered by default job site",
                "Search, sync status badges, pull-to-sync (employees scope)",
                "Registration Step 1 only (details + dept picker from GET /departments)",
                "Employee detail view (profile, sync info)",
                "Employee edit with admin code gate stub (full gate Phase 8)",
                "Delete: local remove + pending deletion queue",
                "Dashboard skeleton (counts from local data)",
            ],
            "api": ["GET/POST /employees", "PUT/DELETE /employees/{id}", "GET /departments"],
            "storage": ["Room Employee table", "pendingEmployeeDeletions", "departmentCatalogJSON"],
            "exit": "Full roster CRUD except face; sync pushes/pulls employee records",
        },
        {
            "title": "Phase 5 — DTR & Attendance Records",
            "goal": "Daily time record grid and attendance sync (no scanner yet).",
            "deliverables": [
                "DTR tab: date picker, default site header, employee rows",
                "Show In/Out times, corrected badge, placeholder for punch photos",
                "Manual test: insert attendance rows locally",
                "Attendance sync: push pending, pull by date, reconcile voids",
                "Pull-to-sync with attendance scope + dtrFocusDate",
                "Dashboard coverage metrics from attendance data",
            ],
            "api": ["GET/POST /attendance (start_date, end_date, updated_since)"],
            "storage": ["Room Attendance", "AttendancePhotoStore file dir"],
            "exit": "DTR grid accurate for default site; syncs punches with server",
        },
        {
            "title": "Phase 6 — Sync Engine (Full Pipeline)",
            "goal": "Complete offline-first sync matching iOS behavior.",
            "deliverables": [
                "SyncService: full push order (device → sites → depts → employees → children → attendance)",
                "SyncMode: quick vs full; SyncScope: all / employees / attendance",
                "Auto-sync foreground timer (3–60 min configurable)",
                "WorkManager periodic background sync",
                "Pending count formula across all entity types",
                "Delta roster sync via lastFullRosterSyncAt",
                "Wi-Fi-only gate for large JPEG uploads",
                "Device registration on every sync (POST /devices)",
            ],
            "api": ["All endpoints from Section 4"],
            "storage": ["All Room tables + pending queues + device registration keys"],
            "exit": "Reliable bidirectional sync; dashboard pending badge accurate",
        },
        {
            "title": "Phase 7 — Face Enrollment & Recognition",
            "goal": "Registration Step 2–3 and attendance scanner.",
            "deliverables": [
                "CameraX preview + face detection (ML Kit or MediaPipe)",
                "Registration Step 2: ID document camera",
                "Registration Step 3: 5-pose enrollment, blink liveness",
                "AdaFace TFLite/ONNX: 512-d embeddings per pose, AES-GCM encrypt",
                "Enrollment JPEG storage per pose",
                "Scanner tab: Time In/Out flow, multi-frame consensus match",
                "Anti-spoof: MiniFASNet TFLite + presentation heuristics",
                "Configurable liveness presets (Settings → Scanner)",
                "Punch photo capture on successful match",
            ],
            "api": ["POST /face-embeddings", "POST /face-enrollment-photos", "POST /employee-id-documents", "POST /attendance"],
            "storage": ["FaceEmbedding, FaceEnrollmentPhoto, EmployeeIdDocument entities + file dirs"],
            "exit": "End-to-end register + scan attendance with face match",
        },
        {
            "title": "Phase 8 — Security Gates & Device Admin Code",
            "goal": "Geofence, supervisor PIN, clock integrity, admin code UI.",
            "deliverables": [
                "PasscodeKeypadView: 6-dot iPhone-style UI",
                "Admin code gates: default site, GPS toggle, job site save, employee edit",
                "POST /devices/verify-admin-code integration",
                "Multi-user assigned operators display from device sync",
                "Supervisor PIN: set/clear, required before punch",
                "Geofence: FusedLocationProvider, block scanner until at default site",
                "Per-employee site validation on punch",
                "ClockIntegrityGuard: server time drift + offline jump detection",
                "DeviceBlockedView + DeviceAccessGuard for SaaS remote disable",
                "X-Device-Local-Id header on all authenticated API calls",
            ],
            "api": ["POST /devices/verify-admin-code", "GET/POST /devices (is_blocked)", "GET /health (server time)"],
            "storage": ["deviceLocalId, assignedUsersJSON, adminCodeRequired, deviceIsBlocked, supervisorPINHash"],
            "exit": "Parity with iOS protected actions, punch validation, and remote block",
        },
        {
            "title": "Phase 9 — Recovery, Diagnostics & Polish",
            "goal": "Production readiness and admin recovery tools.",
            "deliverables": [
                "Restore from cloud backup",
                "Test restore (wipe local + full download)",
                "Advanced diagnostics screen",
                "Debug camera frames toggle",
                "Error handling UX for all sync failures",
                "Pull-to-sync on Settings, Employees, DTR, Dashboard, Job Sites",
                "Notification on sync completion (optional)",
            ],
            "api": ["Full sync pull all entities"],
            "storage": ["Restore clears and repopulates Room + files"],
            "exit": "Feature parity with ConsTrakr iOS for field operations",
        },
    ]

    for i, phase in enumerate(phases, 1):
        doc.add_heading(phase["title"], level=2)
        doc.add_paragraph(f"Goal: {phase['goal']}")
        doc.add_paragraph("Deliverables:", style="List Bullet")
        for d in phase["deliverables"]:
            doc.add_paragraph(d, style="List Bullet 2")
        doc.add_paragraph(f"API: {', '.join(phase['api'])}")
        doc.add_paragraph(f"Storage: {', '.join(phase['storage'])}")
        p = doc.add_paragraph()
        r = p.add_run(f"Exit criteria: {phase['exit']}")
        r.italic = True
        doc.add_paragraph()

    doc.add_page_break()

    # ── 6. Mermaid ──
    doc.add_heading("6. Architecture Diagrams (Mermaid)", level=1)
    doc.add_paragraph(
        "Copy each diagram into https://mermaid.live to render. "
        "These diagrams document system architecture for the Android rebuild."
    )

    add_mermaid_block(doc, "6.1 System Context", """
flowchart TB
    subgraph Web["IMS People (Web)"]
        W1[Employees / DTR / Payroll]
        W2[Job Sites / Devices]
        W3[Void Passcode Setup]
    end
    subgraph API["Constrakr API /constrakr-api/"]
        A1[JWT Auth]
        A2[Sync Endpoints]
    end
    subgraph App["ConsTrakr App (Android)"]
        M1[Scanner / DTR]
        M2[Offline Room DB]
        M3[Sync Queue]
    end
    W1 --> A2
    W2 --> A2
    W3 -.->|admin code| M1
    M3 <-->|Bearer JWT| A2
    M2 <--> M3
    M1 --> M2
""")

    add_mermaid_block(doc, "6.2 App Navigation", """
flowchart TD
    Root[App Root] --> Tabs[Bottom Navigation]
    Tabs --> Dash[Dashboard]
    Tabs --> Emp[Employees]
    Tabs --> Scan[Scanner]
    Tabs --> DTR[Time Record]
    Tabs --> More[More]
    Emp --> Reg[Registration Wizard]
    Emp --> Detail[Employee Detail]
    Detail --> Edit[Employee Edit + Admin Code]
    More --> Sites[Job Sites List]
    Sites --> Editor[Job Site Editor + Admin Code]
    More --> Settings[Settings]
    Settings --> ScannerCfg[Scanner Settings]
    Settings --> SiteGPS[Job Sites and GPS]
    Settings --> SupPIN[Supervisor PIN]
    Settings --> Advanced[Advanced / Restore]
""")

    add_mermaid_block(doc, "6.3 Sync Pipeline", """
sequenceDiagram
    participant UI as UI / Pull-to-Sync
    participant SQ as SyncQueue
    participant SS as SyncService
    participant API as Constrakr API
    participant DB as Room DB

    UI->>SQ: syncNow(mode, scope)
    SQ->>SS: performPushSync()
    SS->>API: POST /devices
    SS->>API: Job sites push/pull
    SS->>API: GET /departments
    SS->>DB: Pending deletions
    SS->>API: GET /employees updated_since
    SS->>DB: Merge remote roster
    SS->>API: POST employees/embeddings/photos/id/attendance
    SS->>API: GET /attendance reconcile voids
    SS->>DB: Update sync status
    SS-->>SQ: Summary + errors
    SQ-->>UI: Refresh UI state
""")

    add_mermaid_block(doc, "6.4 Attendance Scanner Flow", """
flowchart TD
    Start([Open Scanner]) --> Geo{Geofence enabled?}
    Geo -->|Yes| AtSite{At default site?}
    AtSite -->|No| Block[Block UI + Recheck]
    AtSite -->|Yes| Pick[Pick Time In or Out]
    Geo -->|No| Pick
    Pick --> Confirm[Confirm Alert]
    Confirm --> PIN{Supervisor PIN?}
    PIN -->|Yes| EnterPIN[Enter PIN]
    PIN -->|No| Cam[Open Camera]
    EnterPIN --> Cam
    Cam --> Live[Liveness: Blink + Poses + 3D]
    Live --> Spoof[Anti-Spoof Check]
    Spoof --> Match[AdaFace Match 2-frame consensus]
    Match --> Clock[Clock Integrity]
    Clock --> SiteGate[Employee Site GPS Gate]
    SiteGate --> Dup{Duplicate punch?}
    Dup -->|Yes| Already[Already Recorded]
    Dup -->|No| Save[Save Attendance + Photo]
    Save --> Sync[Auto Sync if Online]
""")

    add_mermaid_block(doc, "6.5 Admin Code Gate Flow", """
sequenceDiagram
    participant User
    participant App
    participant API as verify-admin-code
    participant Web as IMS Device Assignment

    Web->>App: assigned_users + adminCodeRequired on sync
    User->>App: Change default site / GPS / Edit employee
    App->>App: ensureChangeAllowed()
    App->>User: PasscodeKeypad 6 digits
    User->>App: Enter code
    App->>API: POST local_id + passcode
    API->>API: Match any assigned user void passcode
    API-->>App: valid true/false
    App->>App: Apply protected change
""")

    add_mermaid_block(doc, "6.6 Android Phase Timeline", """
gantt
    title ConsTrakr Android Rebuild
    dateFormat YYYY-MM-DD
    section Foundation
    Phase 1 Foundation           :p1, 2026-01-01, 14d
    Phase 2 Settings and Auth    :p2, after p1, 10d
    section Core Data
    Phase 3 Job Sites            :p3, after p2, 14d
    Phase 4 Employees            :p4, after p3, 21d
    Phase 5 DTR                  :p5, after p4, 14d
    Phase 6 Sync Engine          :p6, after p5, 21d
    section Face and Security
    Phase 7 Face ML              :p7, after p6, 35d
    Phase 8 Security Gates         :p8, after p7, 14d
    Phase 9 Polish               :p9, after p8, 14d
""")

    add_mermaid_block(doc, "6.7 Local Data Model (Room)", """
erDiagram
    Employee ||--o{ Attendance : has
    Employee ||--o{ FaceEmbedding : has
    Employee ||--o{ FaceEnrollmentPhoto : has
    Employee ||--o| EmployeeIdDocument : has
    Employee {
        uuid localId PK
        string serverId
        string employeeCode
        string firstName
        string lastName
        uuid assignedSiteId
        string syncStatus
    }
    Attendance {
        uuid localId PK
        string serverId
        uuid employeeId
        string checkType
        datetime timestamp
        string syncStatus
    }
    FaceEmbedding {
        uuid id PK
        uuid employeeLocalId
        string pose
        blob encryptedValues
    }
""")

    add_mermaid_block(doc, "6.8 Remote Device Block Flow (SaaS)", """
sequenceDiagram
    participant Admin as Web Admin
    participant API as Constrakr API
    participant App as ConsTrakr App

    Admin->>API: Block device + optional reason
    App->>API: POST /devices (heartbeat) or any sync
    Note over App,API: Header X-Device-Local-Id
    API-->>App: is_blocked true + reason
    App->>App: DeviceBlockedView lock screen
    App--xApp: Scanner / CRUD / upload stopped
    Admin->>API: Unblock device
    App->>API: Check again / sync
    API-->>App: is_blocked false
    App->>App: Dismiss lock, resume normal use
""")

    doc.add_page_break()

    # ── 7. Scope split ──
    doc.add_heading("7. Web-Only vs App Scope", level=1)
    doc.add_heading("Web-only (not required in Android field app)", level=2)
    for b in [
        "Payroll runs, payslips, bank disbursement, 13th month, govt remittance",
        "OT approval workflow",
        "Holiday and leave management",
        "Manual DTR correction (results sync down to app)",
        "Employee payroll setup (allowances, deductions, rate overrides)",
        "Reports hub and CSV exports",
        "Support tools (enrollment gaps, sync integrity)",
        "Void punch/day (web action; voids sync to app via attendance pull)",
        "Device user assignment (web configures; app consumes via device API)",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("App-only (field operations)", level=2)
    for b in [
        "Face enrollment and live attendance scanner",
        "Offline-first roster and punch storage",
        "Geofence and supervisor PIN at punch time",
        "Default job site selection with admin code",
        "Device registration and local device name",
        "Cloud restore on replacement device",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_page_break()

    # ── 8. Tech mapping ──
    doc.add_heading("8. Android Technology Mapping", level=1)
    add_table(doc, ["iOS", "Android"], [
        ["SwiftUI", "Jetpack Compose"],
        ["SwiftData", "Room"],
        ["UserDefaults + JSON", "SharedPreferences / DataStore"],
        ["Keychain", "EncryptedSharedPreferences / Keystore"],
        ["BGAppRefreshTask", "WorkManager"],
        ["Core ML AdaFace/MiniFAS", "TFLite or ONNX Runtime"],
        ["Vision face landmarks", "ML Kit Face Detection / MediaPipe"],
        ["CLLocationManager", "FusedLocationProviderClient"],
        ["AVFoundation camera", "CameraX"],
        ["AES-GCM encryption", "Android Cipher AES/GCM"],
    ])

    doc.add_heading("Critical Business Rules to Preserve", level=2)
    rules = [
        "Default job site drives Dashboard, Employee list, DTR, and scanner operating context",
        "Employee assigned site overrides default for punch geofence validation",
        "Offline-first: always write locally as pending; sync when online",
        "One Time In + one Time Out per employee per calendar day",
        "Face enrollment = 5 poses; scanner liveness = configurable subset + blink + optional 3D",
        "Admin code = server-verified 6-digit; Supervisor PIN = local optional gate",
        "Employee delete = local wipe + server soft-delete queue",
        "DTR sync scope = attendance only with date focus",
        "Embeddings encrypted at rest and on wire (AES-GCM)",
        "Clock integrity check before every punch",
        "Remote block: web is_blocked stops scan, CRUD, and sync; enforced server-side via X-Device-Local-Id",
    ]
    for r in rules:
        doc.add_paragraph(r, style="List Number")

    return doc


def main() -> None:
    out = "/Users/mariagemmalusterio/Desktop/ConsTrakr/ConsTrakr-Android-Rebuild-Spec.docx"
    repo_out = "/Users/mariagemmalusterio/Desktop/ConsTrakr/ConsTrakr/docs/ConsTrakr-Android-Rebuild-Spec.docx"
    doc = build_document()
    doc.save(out)
    import os
    os.makedirs(os.path.dirname(repo_out), exist_ok=True)
    doc.save(repo_out)
    print(f"Written: {out}")
    print(f"Written: {repo_out}")


if __name__ == "__main__":
    main()
